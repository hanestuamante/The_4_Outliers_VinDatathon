# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copy2
import gc
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[3]
DOC_PATH = ROOT / "report draft 304.docx"
BACKUP_PATH = ROOT / "report draft 304.before_eda_rewrite.docx"
FIG_DIR = ROOT / "VinDatathon_the-4-Outliers" / "outputs" / "report_revision"
FIG_A = FIG_DIR / "fig_insight_a_promo_capital_trap.png"
FIG_B = FIG_DIR / "fig_insight_b_portfolio_drag.png"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
WORD_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def compact(paragraph, after=1.5, before=0):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def prune_unused_document_media(docx_path):
    """Remove stale image relationships left over from the original draft."""
    rels_path = "word/_rels/document.xml.rels"
    doc_xml_path = "word/document.xml"
    rid_attr = f"{{{WORD_REL_NS}}}embed"

    with ZipFile(docx_path, "r") as zin:
        doc_root = ET.fromstring(zin.read(doc_xml_path))
        used_rids = {node.attrib[rid_attr] for node in doc_root.iter() if rid_attr in node.attrib}

        rels_root = ET.fromstring(zin.read(rels_path))
        kept_media = set()
        for rel in list(rels_root):
            target = rel.attrib.get("Target", "")
            rid = rel.attrib.get("Id")
            rel_type = rel.attrib.get("Type")
            if rel_type == IMAGE_REL_TYPE and target.startswith("media/") and rid not in used_rids:
                rels_root.remove(rel)
            elif rel_type == IMAGE_REL_TYPE and target.startswith("media/"):
                kept_media.add(f"word/{target}")

        ET.register_namespace("", REL_NS)
        new_rels = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path.with_suffix(".tmp.docx")

        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == rels_path:
                    zout.writestr(item, new_rels)
                    continue
                if item.filename.startswith("word/media/") and item.filename not in kept_media:
                    continue
                zout.writestr(item, zin.read(item.filename))

    try:
        tmp_path.replace(docx_path)
    except PermissionError:
        tmp_path.unlink(missing_ok=True)
        print(f"Skipped media pruning because the document is locked: {docx_path}")


def main():
    if not BACKUP_PATH.exists():
        raise FileNotFoundError(f"Missing original backup: {BACKUP_PATH}")
    if not FIG_A.exists() or not FIG_B.exists():
        raise FileNotFoundError("Expected revised figures are missing.")

    copy2(BACKUP_PATH, DOC_PATH)
    doc = Document(DOC_PATH)

    doc.paragraphs[8].text = (
        "Báo cáo phân tích dữ liệu thương mại điện tử thời trang Việt Nam giai đoạn 2012-2022 nhằm kết nối EDA kinh doanh với bài toán dự báo daily Revenue/COGS. "
        "Nhóm xây dựng revenue lineage để phân biệt gross booked revenue, net paid after discount và realized cash after refund, từ đó tránh trộn lẫn forecasting target với KPI kinh tế sau discount/refund. "
        "EDA tập trung vào hai cơ chế xói mòn có thể can thiệp: Promo Capital Trap (chiết khấu tạo margin âm, mua lại khách cũ và bị dùng sai mùa) và Portfolio Drag (long-tail SKU cùng wrong-size returns phá lợi nhuận gộp). "
        "Các tín hiệu này được chuyển hóa thành feature dự báo như seasonality, lag/rolling, holiday window, traffic, promotion intensity, inventory/stockout và return leakage indicators."
    )

    for p in list(doc.paragraphs)[15:]:
        remove_paragraph(p)

    def add_para(text="", style="Normal", after=1.5, before=0):
        p = doc.add_paragraph(style=style)
        compact(p, after=after, before=before)
        if text:
            p.add_run(text)
        return p

    def add_labeled(label, text):
        p = add_para()
        r = p.add_run(label)
        r.bold = True
        p.add_run(" " + text)
        return p

    def add_figure(path, caption):
        p = doc.add_paragraph()
        compact(p, after=0.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.15))
        cap = add_para(caption, after=2.0)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(8)

    add_para(
        "Từ 5 trụ cột, nhóm chọn 2 cơ chế có tín hiệu dữ liệu rõ nhất và actionability cao nhất để trình bày trong phạm vi 1.75 trang:",
        after=0.5,
    )
    add_para("Insight I: Promo Capital Trap - khuyến mãi đang bù lỗ cho khách cũ và dùng sai thời điểm nhu cầu.", after=0.5)
    add_para("Insight II: Portfolio Drag - long-tail SKU và wrong-size returns đang ăn mòn lợi nhuận từ nhóm hàng chủ lực.", after=0.5)
    add_para(
        "Điểm chung: doanh thu đặt hàng cao chưa đồng nghĩa với giá trị giữ lại cao; cần quản trị vốn promo, tồn kho và chất lượng danh mục cùng lúc.",
        after=3.0,
    )

    add_para("Phân tích khám phá dữ liệu", style="Heading 1", after=3.0)
    add_para(
        "Gross booked revenue đạt 16.43B VND, nhưng sau 749.6M VND discount và 510.6M VND refund, giá trị giữ lại chỉ còn khoảng 15.17B VND. "
        "Thay vì dàn trải nhiều pattern, phần EDA tập trung vào hai cơ chế có thể can thiệp ngay: vốn promo bị đốt sai mục tiêu và danh mục/return kéo tụt GP.",
        after=3.0,
    )

    add_para("Insight A - Promo Capital Trap: sai người, sai thời điểm", style="Heading 2", after=1.0, before=2.0)
    add_labeled(
        "Descriptive.",
        "Promo tạo net fulfilled margin -14.4% so với +20.0% ở non-promo, tức gap 34.4pp. Trên cùng fulfilled basis, 514.7M/671.9M VND = 76.6% discount burn nằm ở off-peak, nơi revenue/day chỉ bằng khoảng 55% mùa cao điểm.",
    )
    add_labeled(
        "Diagnostic.",
        "95.7% promo net sales đến từ existing customers; H1 acquisition bị bác bỏ vì r(buyers, promo rate) = +0.048. H2 growth cũng không được hỗ trợ khi r(promo rate, net revenue) = -0.212. Nghịch lý thời điểm làm vấn đề nặng hơn: Apr-Jun có revenue/day 1.81x và sessions/day 1.68x nhưng promo rate chỉ 23.8% so với 43.4% off-peak; cùng lúc peak stockout lên tới 10,537 ngày ở Streetwear và 5,586 ngày ở Outdoor.",
    )
    add_figure(
        FIG_A,
        "Figure 1. Promo Capital Trap: margin âm và promo lệch mùa; 76.6% dùng fulfilled discount basis (514.7M/671.9M VND).",
    )
    add_labeled(
        "Predictive.",
        "Redirect 30% off-peak discount tương đương 154.4M VND sang pre-peak inventory cho Streetwear/Outdoor cho GP recovery upper bound khoảng 50.0M VND. Cắt 30% weak-promo exposure có thể tránh khoảng 202M VND discount burn; full margin-gap recovery là upper bound khoảng 433M VND trước điều chỉnh volume loss.",
    )
    add_labeled(
        "Prescriptive.",
        "Ưu tiên ngay (sprint 1): đặt guardrail net margin >= 0 cho campaign rollout và chuyển voucher sang first-purchase-only; chạy regional A/B holdout 4 tuần. Trước Apr-Jun 6-8 tuần, tái phân bổ 154.4M VND vào stock readiness cho Streetwear/Outdoor. KPI: new-customer promo share >15%, net fulfilled margin 9.7% -> 14-15%, peak stockout rate <50% (baseline Streetwear: 68.3%).",
    )

    add_para("Insight B - Portfolio Drag: sản phẩm sai và size-guide sai đang phá GP", style="Heading 2", after=1.0, before=2.0)
    add_labeled(
        "Descriptive.",
        "Net fulfilled sales đạt 14.052B VND với margin 9.69%, nhưng GP tập trung cực đoan: top 20% SKU tạo 121.9% net GP trong khi bottom 30% phá hủy 25.8%. 521 negative-GP SKUs gây 350.9M VND historical loss; wrong_size là lý do return #1 với 13,967 returns, 34.97% reasons và 176.7M VND refund trong pool 2012-2022.",
    )
    add_labeled(
        "Diagnostic.",
        "Return rate theo size S/M/L/XL gần như phẳng 3.75%-3.86%, nên lỗi nằm ở size guide/PDP expectation chứ không phải một size cụ thể. Trong long tail, 231 low-volume negative-GP SKUs là quarantine candidates, gây 13.4M VND loss; đồng thời Streetwear chiếm 80.1% revenue nhưng margin 9.29%, còn GenZ chỉ 2.10% revenue nhưng margin 15.41%.",
    )
    add_figure(
        FIG_B,
        "Figure 2. Portfolio Drag: GP lệ thuộc top SKU, còn wrong_size là lỗi hệ thống vì return rate phẳng theo size.",
    )
    add_labeled(
        "Predictive.",
        "Giảm 40-60% wrong_size returns có thể bảo toàn 70.7-106.0M VND trong historical comparable refund pool. Quarantine 231 SKU giúp chặn 13.4M VND drag, còn loại 656 ghost_fake SKUs khỏi analytics giúp mô hình và dashboard không học nhiễu từ mã không vận hành.",
    )
    add_labeled(
        "Prescriptive.",
        "Thử nghiệm PDP size filter/guide trong 6 tuần, quarantine 231 SKU trong 30 ngày, và làm sạch ghost_fake SKU khỏi analytics layer. KPI: return rate 6.22% -> 5.1%, wrong-size share <20%, negative-GP SKU count giảm 25% trong một quý.",
    )

    add_para(
        "Closing bridge. Hai cơ chế trên độc lập nhưng cùng hướng: chuyển discount phản ứng thành inventory readiness và cắt SKU/return drag. Upside cần đọc như proxy/upper-bound, không cộng thẳng thành realized P&L: khoảng 50.0M VND GP từ stockout-capture proxy, 70.7-106.0M VND wrong-size refund pool, và 202M VND discount exposure avoided. Forecast model nên encode seasonal demand, promo intensity, stockout/inventory friction và return leakage như leading signals.",
        after=2.0,
    )
    add_para(
        "Metric note: 76.6% dùng cùng fulfilled denominator (514.7M off-peak fulfilled discount / 671.9M fulfilled promo discount). Các recovery estimate đều là scenario hoặc upper-bound trước khi kiểm chứng bằng holdout/A-B test.",
        after=0.5,
    )

    for p in doc.paragraphs:
        if p.style.name == "Normal":
            compact(p, after=p.paragraph_format.space_after.pt if p.paragraph_format.space_after else 1.5)

    doc.save(DOC_PATH)
    del doc
    gc.collect()
    prune_unused_document_media(DOC_PATH)
    print(f"Rewrote {DOC_PATH}")


if __name__ == "__main__":
    main()
