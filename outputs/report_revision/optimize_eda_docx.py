# -*- coding: utf-8 -*-
"""Rewrite the EDA section with audit-safe wording and refreshed figures.

The script rebuilds `report draft 304.docx` from the original pre-EDA-rewrite
backup, then inserts the optimized two-insight EDA narrative. It keeps the
metric basis explicit so the report can be traced back to notebooks 3, 4, and 6.
"""

from pathlib import Path
from shutil import copy2
import gc
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[3]
DOC_PATH = ROOT / "report draft 304.docx"
OPT_PATH = ROOT / "report draft 304.optimized_eda.docx"
ORIGINAL_BACKUP_PATH = ROOT / "report draft 304.before_eda_rewrite.docx"
MAX_BACKUP_PATH = ROOT / "report draft 304.before_max_eda_optimization.docx"

CHART_OUT = (
    ROOT
    / "VinDatathon_the-4-Outliers"
    / "5 góc nhìn"
    / "report_chart_source"
    / "outputs"
)
FIG_A = CHART_OUT / "fig_report_insight_a_promo_capital_trap.png"
FIG_B = CHART_OUT / "fig_report_insight_b_portfolio_drag.png"

REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
WORD_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
IMAGE_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def compact(paragraph, after=1.2, before=0):
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def prune_unused_document_media(docx_path):
    """Remove stale image relationships left over from the original draft."""
    rels_path = "word/_rels/document.xml.rels"
    doc_xml_path = "word/document.xml"
    rid_attr = f"{{{WORD_REL_NS}}}embed"

    with ZipFile(docx_path, "r") as zin:
        doc_root = ET.fromstring(zin.read(doc_xml_path))
        used_rids = {node.attrib[rid_attr] for node in doc_root.iter() if rid_attr in node.attrib}

        rels_root = ET.fromstring(zin.read(rels_path))
        kept_media = set()
        for rel in list(rels_root):
            target = rel.attrib.get("Target", "")
            rid = rel.attrib.get("Id")
            rel_type = rel.attrib.get("Type")
            if rel_type == IMAGE_REL_TYPE and target.startswith("media/") and rid not in used_rids:
                rels_root.remove(rel)
            elif rel_type == IMAGE_REL_TYPE and target.startswith("media/"):
                kept_media.add(f"word/{target}")

        ET.register_namespace("", REL_NS)
        new_rels = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
        tmp_path = docx_path.with_suffix(".tmp.docx")

        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == rels_path:
                    zout.writestr(item, new_rels)
                    continue
                if item.filename.startswith("word/media/") and item.filename not in kept_media:
                    continue
                zout.writestr(item, zin.read(item.filename))

    try:
        tmp_path.replace(docx_path)
    except PermissionError:
        tmp_path.unlink(missing_ok=True)
        print(f"Skipped media pruning because the document is locked: {docx_path}")


def set_run_font(paragraph, size_pt=None, italic=None, bold=None):
    for run in paragraph.runs:
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if italic is not None:
            run.italic = italic
        if bold is not None:
            run.bold = bold


def main():
    if not ORIGINAL_BACKUP_PATH.exists():
        raise FileNotFoundError(f"Missing original backup: {ORIGINAL_BACKUP_PATH}")
    if not FIG_A.exists() or not FIG_B.exists():
        raise FileNotFoundError("Expected report figures are missing. Run make_report_charts.py first.")

    if DOC_PATH.exists() and not MAX_BACKUP_PATH.exists():
        copy2(DOC_PATH, MAX_BACKUP_PATH)

    copy2(ORIGINAL_BACKUP_PATH, DOC_PATH)
    doc = Document(DOC_PATH)

    doc.paragraphs[8].text = (
        "Báo cáo phân tích dữ liệu thương mại điện tử thời trang Việt Nam giai đoạn 2012-2022 "
        "kết nối EDA kinh doanh với bài toán dự báo daily Revenue/COGS. Nhóm xây dựng revenue "
        "lineage để tách gross booked revenue (forecast target), net paid after discount "
        "(unit economics) và refund/return leakage (diagnostic pool), tránh cộng lẫn các basis "
        "khác nhau thành P&L nếu chưa có holdout. EDA tập trung vào hai cơ chế có thể can thiệp: "
        "Promo Capital Trap (margin âm, mua lại khách cũ, dùng sai mùa) và Portfolio Drag "
        "(long-tail SKU cùng wrong-size returns ăn mòn GP). Các tín hiệu này được chuyển hóa "
        "thành feature dự báo như seasonality, promo intensity, inventory/stockout friction và "
        "return leakage indicators."
    )

    doc.paragraphs[12].text = (
        "Để đảm bảo tính nhất quán, nhóm dùng ba basis tách biệt. Gross booked revenue được dùng "
        "cho forecasting vì khớp trực tiếp với target Revenue/COGS trong sales.csv. Net paid/net "
        "fulfilled sau discount được dùng cho unit economics của promo và SKU. Refund/return được "
        "đọc như leakage pool riêng để chẩn đoán chất lượng giữ lại giá trị; các scenario trong "
        "EDA là proxy hoặc upper-bound trước khi kiểm chứng bằng holdout/A-B test."
    )

    for p in list(doc.paragraphs)[15:]:
        remove_paragraph(p)

    def add_para(text="", style="Normal", after=1.2, before=0, size=None):
        p = doc.add_paragraph(style=style)
        compact(p, after=after, before=before)
        if text:
            p.add_run(text)
        if size is not None:
            set_run_font(p, size_pt=size)
        return p

    def add_labeled(label, text):
        p = add_para(after=1.0)
        r = p.add_run(label)
        r.bold = True
        p.add_run(" " + text)
        return p

    def add_figure(path, caption):
        p = doc.add_paragraph()
        compact(p, after=0.3)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        section = doc.sections[0]
        content_width = section.page_width - section.left_margin - section.right_margin
        run.add_picture(str(path), width=content_width)

        cap = add_para(caption, after=1.4, size=8)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cap, size_pt=8, italic=True)

    add_para(
        "Từ 5 trụ cột, nhóm chỉ giữ 2 cơ chế có tín hiệu dữ liệu rõ nhất và actionability cao "
        "nhất trong giới hạn 1.75 trang:",
        after=0.35,
    )
    add_para(
        "Insight I: Promo Capital Trap - khuyến mãi đang bù lỗ cho khách cũ và dùng sai thời "
        "điểm nhu cầu.",
        after=0.35,
    )
    add_para(
        "Insight II: Portfolio Drag - long-tail SKU và wrong-size returns đang ăn mòn lợi nhuận "
        "từ nhóm hàng chủ lực.",
        after=0.35,
    )
    add_para(
        "Thesis chung: doanh thu đặt hàng cao chưa đồng nghĩa với giá trị giữ lại cao; muốn dự báo "
        "và vận hành tốt hơn cần quản trị vốn promo, tồn kho và chất lượng danh mục cùng lúc.",
        after=2.0,
    )

    add_para("Phân tích khám phá dữ liệu", style="Heading 1", after=2.0)
    add_para(
        "Gross booked revenue đạt 16.43B VND; sau 749.6M VND discount và 510.6M VND refund "
        "từ bảng returns, realized-cash proxy còn khoảng 15.17B VND. Hai insight dưới đây tách "
        "basis: promo/SKU margin dùng net fulfilled after discount; return dùng historical refund "
        "pool 2012-2022; recovery estimate chỉ là scenario/proxy trước holdout.",
        after=2.0,
    )

    add_para("Insight A - Promo Capital Trap: sai người, sai thời điểm", style="Heading 2", after=0.8, before=1.2)
    add_labeled(
        "Descriptive/Diagnostic.",
        "Promo có net fulfilled margin -14.4% so với +20.0% non-promo (gap 34.4pp); 76.6% "
        "fulfilled discount burn nằm off-peak (514.7M/671.9M VND), dù off-peak revenue/day chỉ "
        "bằng khoảng 55% peak. Trong 10 campaign gần nhất, 95.7% promo net sales đến từ existing "
        "customers; r(buyers,promo)=+0.048 và r(revenue,promo)=-0.212 không ủng hộ giả thuyết "
        "acquisition/volume protection. Trên K3 seasonality basis, Apr-Jun có revenue/day 1.81x "
        "và sessions/day 1.68x ở promo rate thấp hơn (23.8% vs 43.4%), trong khi inventory ghi "
        "nhận peak stockout-days: Streetwear 10,537 và Outdoor 5,586.",
    )
    add_figure(
        FIG_A,
        "Figure 1. Promo Capital Trap: evidence pattern cho margin âm, existing-heavy promos và "
        "off-peak discount timing. Seasonality dùng K3 gross/monthly basis; discount share dùng "
        "fulfilled basis (514.7M/671.9M VND).",
    )
    add_labeled(
        "Predictive/Prescriptive.",
        "Xem off-peak promo như reallocation pool, không phải uplift đã chứng minh: 30% redirect = "
        "154.4M VND funding envelope; Notebook 6 ước tính targeted stockout-capture upper bound "
        "~50.0M VND GP; Notebook 3 ước tính weak-promo discount avoided ~202M VND. Sprint 1 "
        "guardrails: net margin >= 0, first-purchase voucher, regional A/B holdout. Pre-peak "
        "6-8 weeks, allocate only to forecast-supported, positive-margin SKU/category with observed "
        "stockout and high GP per inventory VND. KPI test: new-customer promo share >15%, margin "
        "9.7% -> 14-15%, peak stockout SKU-month rate <50% (baseline overall 68.3%; Streetwear 67.8%).",
    )

    add_para("Insight B - Portfolio Drag: long-tail SKU và wrong-size returns phá GP", style="Heading 2", after=0.8, before=1.2)
    add_labeled(
        "Descriptive/Diagnostic.",
        "Net fulfilled sales đạt 14.052B VND với margin 9.69%, nhưng GP tập trung cực đoan: top "
        "20% master SKUs tạo 121.9% net GP, còn bottom 30% phá hủy 25.8%. 521 negative-GP SKUs "
        "gây 350.9M VND historical loss. wrong_size là return reason #1 (13,967 return rows, "
        "34.97%, 176.7M VND refund, 2012-2022). Vì wrong-size share theo S/M/L/XL gần phẳng "
        "34.68%-35.47% và return quantity rate cũng chỉ dao động 3.75%-3.86%, evidence phù hợp "
        "với systemic fit-guidance/PDP issue hơn là một size lỗi. Long tail có 231 quarantine "
        "candidates, gây 13.4M VND loss.",
    )
    add_figure(
        FIG_B,
        "Figure 2. Portfolio Drag: GP phụ thuộc top master SKUs, wrong_size là refund pool lớn nhất, "
        "và wrong-size share gần phẳng theo size; các loss/refund là historical 2012-2022, không annualized.",
    )
    add_labeled(
        "Predictive/Prescriptive.",
        "Giảm 40-60% wrong_size returns bảo toàn 70.7-106.0M VND của historical comparable refund "
        "pool, không annualized. Chạy PDP size-guide/filter test 6 tuần với control group; quarantine "
        "231 SKU trong 30 ngày trước hard delist; loại 656 ghost_fake SKUs khỏi analytics. KPI test: "
        "return rate 6.22% -> khoảng 5.1%, wrong-size share <25% first milestone và <20% stretch, "
        "negative-GP SKU count giảm 25% trong một quý.",
    )

    add_para(
        "Closing bridge. Không cộng thẳng upside thành realized P&L: 50.0M VND GP, 70.7-106.0M "
        "VND refund pool và 202M VND discount avoided là các basis khác nhau, cần holdout. Forecast "
        "model nên encode seasonal demand, promo intensity, stockout/inventory friction và return "
        "leakage như leading signals.",
        after=1.2,
    )
    add_para(
        "Metric note: fulfilled/committed basis loại cancelled/created và giữ returned cho demand "
        "planning; refund được đọc riêng trong return leakage. 510.6M VND refund lấy từ returns table; "
        "mọi recovery estimate là scenario/proxy/upper-bound trước khi kiểm chứng bằng holdout hoặc A/B test.",
        after=0.2,
        size=8,
    )

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Normal":
            if paragraph.paragraph_format.space_after is None:
                compact(paragraph)

    doc.save(DOC_PATH)
    del doc
    gc.collect()
    prune_unused_document_media(DOC_PATH)
    copy2(DOC_PATH, OPT_PATH)
    print(f"Optimized EDA written to: {DOC_PATH}")
    print(f"Optimized copy written to: {OPT_PATH}")


if __name__ == "__main__":
    main()
